# ══════════════════════════════════════════
# Дополнительные тесты: картинки, стили ячеек,
# colspan, XML-fallback цвета, выравнивание, PDF
# ══════════════════════════════════════════
import io
from pathlib import Path
from unittest.mock import MagicMock, patch

from docx import Document

from app.core import file_parser
from app.core.file_parser import extract_text_from_file

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.file_parser import (
    _extract_images,
    _get_cell_span,
    _get_cell_styles,
    _get_run_color,
    _parse_pdf,
)

# минимальный валидный 1x1 PNG (~70 байт)
_PNG_1x1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
    b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _make_docx(tmp_path, build_fn) -> Path:
    doc = Document()
    build_fn(doc)
    path = tmp_path / "styled.docx"
    doc.save(str(path))
    return path


# --- XML-хелперы для правки ячеек ---
def _set_cell_shading(cell, fill: str):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_borders(cell):
    tcPr = cell._element.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")

    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "16")          # 16//8 = 2px
    top.set(qn("w:color"), "FF0000")
    borders.append(top)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "bad")      # ValueError -> px=1, цвет дефолтный
    borders.append(bottom)

    left = OxmlElement("w:left")
    left.set(qn("w:val"), "none")      # пропускается
    borders.append(left)
    # right отсутствует -> ветка "continue"

    tcPr.append(borders)


def _set_cell_valign(cell, val: str):
    tcPr = cell._element.get_or_add_tcPr()
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), val)
    tcPr.append(valign)


def _set_cell_gridspan(cell, span: int):
    tcPr = cell._element.get_or_add_tcPr()
    gs = OxmlElement("w:gridSpan")
    gs.set(qn("w:val"), str(span))
    tcPr.append(gs)


# ── Картинки в docx (строки 180-192, 310-314) ──
class TestDocxImages:
    def test_inline_image_rendered(self, tmp_path):
        path = _make_docx(
            tmp_path, lambda d: d.add_picture(io.BytesIO(_PNG_1x1))
        )
        result = extract_text_from_file(path)
        assert '<img src="data:image/png;base64,' in result

    def test_images_budget_exceeded_skips_rest(self, tmp_path):
        # лимит почти на нуле -> первая картинка уже превышает бюджет
        path = _make_docx(
            tmp_path, lambda d: d.add_picture(io.BytesIO(_PNG_1x1))
        )
        doc = Document(str(path))
        with patch.object(file_parser, "_MAX_INLINE_IMAGES_BYTES", 1):
            images = _extract_images(doc)
        assert images == {}

    def test_extract_images_handles_broken_rels(self):
        # rels кидает исключение -> ловится, возвращается {}
        fake_doc = MagicMock()
        type(fake_doc.part).rels = property(
            lambda self: (_ for _ in ()).throw(RuntimeError("boom"))
        )
        assert _extract_images(fake_doc) == {}


# ── Выравнивание абзаца (строка 245) ──
class TestParagraphAlignment:
    def test_center_alignment_adds_style(self, tmp_path):
        def build(d):
            p = d.add_paragraph("по центру")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        result = extract_text_from_file(_make_docx(tmp_path, build))
        assert 'style="text-align: center"' in result

    def test_left_alignment_no_style(self, tmp_path):
        def build(d):
            p = d.add_paragraph("слева")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        result = extract_text_from_file(_make_docx(tmp_path, build))
        assert "text-align" not in result


# ── XML-fallback цвета через w:rPr (строки 287-301) ──
class TestRunColorXmlFallback:
    def _run_with_xml_color(self, color_val: str):
        """Создаёт run, где цвет задан только в XML (w:color), без font.color.rgb."""
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("текст")
        rpr = run._element.get_or_add_rPr()
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color_val)
        rpr.append(color_el)
        return run

    def test_color_from_xml(self):
        run = self._run_with_xml_color("00FF00")
        # font.color.rgb тут None -> уходим во второй try (XML)
        assert _get_run_color(run) == "#00FF00"

    def test_color_auto_ignored(self):
        run = self._run_with_xml_color("auto")
        assert _get_run_color(run) is None

    def test_no_rpr_returns_none(self):
        doc = Document()
        p = doc.add_paragraph()
        run = p.add_run("без стилей")
        assert _get_run_color(run) is None


# ── Стили ячеек таблицы (строки 351-384) ──
class TestCellStyles:
    def test_shading_border_valign(self, tmp_path):
        def build(d):
            table = d.add_table(rows=2, cols=1)
            cell = table.cell(1, 0)
            cell.text = "ячейка"
            _set_cell_shading(cell, "FFFF00")
            _set_cell_borders(cell)
            _set_cell_valign(cell, "center")

        result = extract_text_from_file(_make_docx(tmp_path, build))
        assert "background-color: #FFFF00" in result
        assert "border-top: 2px solid #FF0000" in result
        assert "border-bottom: 1px solid #000000" in result  # sz=bad -> 1px
        assert "vertical-align: middle" in result

    def test_shading_auto_skipped(self, tmp_path):
        def build(d):
            table = d.add_table(rows=2, cols=1)
            cell = table.cell(1, 0)
            cell.text = "x"
            _set_cell_shading(cell, "auto")

        result = extract_text_from_file(_make_docx(tmp_path, build))
        assert "background-color" not in result

    def test_no_tcpr_returns_empty(self):
        fake_cell = MagicMock()
        fake_cell._element.find.return_value = None
        assert _get_cell_styles(fake_cell) == ""

    def test_cell_styles_exception_handled(self):
        fake_cell = MagicMock()
        fake_cell._element.find.side_effect = RuntimeError("boom")
        assert _get_cell_styles(fake_cell) == ""


# ── colspan (строки 393-402) ──
class TestCellSpan:
    def test_colspan_rendered(self, tmp_path):
        def build(d):
            table = d.add_table(rows=2, cols=2)
            cell = table.cell(1, 0)
            cell.text = "широкая"
            _set_cell_gridspan(cell, 2)

        result = extract_text_from_file(_make_docx(tmp_path, build))
        assert 'colspan="2"' in result

    def test_gridspan_one_no_attr(self, tmp_path):
        def build(d):
            table = d.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)
            cell.text = "обычная"
            _set_cell_gridspan(cell, 1)

        result = extract_text_from_file(_make_docx(tmp_path, build))
        assert "colspan" not in result

    def test_no_tcpr_no_span(self):
        fake_cell = MagicMock()
        fake_cell._element.find.return_value = None
        assert _get_cell_span(fake_cell) == ""

    def test_span_exception_handled(self):
        fake_cell = MagicMock()
        fake_cell._element.find.side_effect = RuntimeError("boom")
        assert _get_cell_span(fake_cell) == ""


# ── PDF: тело _parse_pdf напрямую (строки 121-124) ──
class TestParsePdfDirect:
    def test_parse_pdf_body(self, tmp_path):
        f = tmp_path / "x.pdf"
        f.write_bytes(b"%PDF-fake")

        fake_page = MagicMock()
        fake_page.get_text.return_value = "Абзац 1\nАбзац 2\n   \n"
        fake_doc = MagicMock()
        fake_doc.__enter__.return_value = [fake_page]
        fake_doc.__exit__.return_value = False
        fake_fitz = MagicMock()
        fake_fitz.open.return_value = fake_doc

        with patch.object(file_parser, "fitz", fake_fitz):
            result = _parse_pdf(f)

        assert result == "<p>Абзац 1</p>\n<p>Абзац 2</p>"
