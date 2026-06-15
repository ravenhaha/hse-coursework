"""Баг №3: парсер привязан к разрешённым форматам.

extract_text_from_file:
    - .txt / .md         → HTML по абзацам;
    - .docx              → HTML (python-docx);
    - .pdf               → HTML, если установлен PyMuPDF;
    - картинки/прочее    → None (парсера нет — это норма);
    - битый файл         → None (исключение проглочено, не пробрасывается).

Тесты не трогают БД и API — только чистая функция парсинга.
"""

import pytest
from docx import Document

from app.core.constants import (
    ALLOWED_MATERIAL_EXTENSIONS,
    MATERIAL_DOC_EXT,
    MATERIAL_IMAGE_EXT,
    MATERIAL_TEXT_EXT,
)
from app.core.file_parser import _PARSERS, extract_text_from_file


# ══════════════════════════════════════════
# Контракт констант ↔ парсеров
# ══════════════════════════════════════════
class TestExtensionsContract:
    @pytest.mark.parametrize("ext", [".rtf", ".xlsx", ".csv"])
    def test_removed_formats_not_allowed(self, ext):
        """Лишние форматы убраны из списка разрешённых (регресс-защита)."""
        assert ext not in ALLOWED_MATERIAL_EXTENSIONS

    def test_text_and_doc_formats_have_parser(self):
        """Каждое текстовое/документное расширение имеет парсер.

        PDF исключаем: его парсер регистрируется только при наличии
        PyMuPDF, иначе он осознанно отсутствует.
        """
        parseable = (MATERIAL_TEXT_EXT | MATERIAL_DOC_EXT) - {".pdf"}
        for ext in parseable:
            assert ext in _PARSERS, f"нет парсера для {ext}"

    def test_images_have_no_parser_by_design(self):
        """Картинки парсера не имеют — текст из них не извлекается (без OCR)."""
        for ext in MATERIAL_IMAGE_EXT:
            assert ext not in _PARSERS


# ══════════════════════════════════════════
# TXT / MD
# ══════════════════════════════════════════
class TestTextParser:
    def test_txt_to_paragraphs(self, tmp_path):
        f = tmp_path / "note.txt"
        f.write_text("Первая строка\nВторая строка\n", encoding="utf-8")

        html = extract_text_from_file(f)
        assert html == "<p>Первая строка</p>\n<p>Вторая строка</p>"

    def test_empty_txt_gives_empty_paragraph(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("   \n\n", encoding="utf-8")
        assert extract_text_from_file(f) == "<p></p>"

    def test_html_is_escaped(self, tmp_path):
        """Текст экранируется — защита от XSS при рендере во фронте."""
        f = tmp_path / "xss.txt"
        f.write_text("<script>alert(1)</script>", encoding="utf-8")

        html = extract_text_from_file(f)
        assert "<script>" not in html
        assert "&lt;script&gt;" in html

    def test_md_uses_same_parser(self, tmp_path):
        f = tmp_path / "doc.md"
        f.write_text("# Заголовок\nтекст", encoding="utf-8")

        html = extract_text_from_file(f)
        # Markdown НЕ рендерится — '#' остаётся как обычный текст
        assert "<p># Заголовок</p>" in html


# ══════════════════════════════════════════
# DOCX
# ══════════════════════════════════════════
class TestDocxParser:
    def test_docx_paragraphs(self, tmp_path):
        doc = Document()
        doc.add_paragraph("Привет, мир")
        doc.add_paragraph("Второй абзац")
        path = tmp_path / "test.docx"
        doc.save(path)

        html = extract_text_from_file(path)
        assert "Привет, мир" in html
        assert "Второй абзац" in html
        assert "<p" in html

    def test_docx_escapes_html(self, tmp_path):
        doc = Document()
        doc.add_paragraph("<b>не тег</b>")
        path = tmp_path / "esc.docx"
        doc.save(path)

        html = extract_text_from_file(path)
        assert "&lt;b&gt;" in html


# ══════════════════════════════════════════
# Неподдерживаемое / битое
# ══════════════════════════════════════════
class TestUnsupportedAndBroken:
    @pytest.mark.parametrize("name", ["pic.jpg", "pic.png", "data.csv", "noext"])
    def test_unsupported_returns_none(self, tmp_path, name):
        """Нет зарегистрированного парсера → None (без исключений)."""
        f = tmp_path / name
        f.write_bytes(b"whatever")
        assert extract_text_from_file(f) is None

    def test_broken_docx_returns_none(self, tmp_path):
        """Битый docx не валит приложение — парсер ловит исключение."""
        f = tmp_path / "broken.docx"
        f.write_bytes(b"this is not a real docx")
        assert extract_text_from_file(f) is None
